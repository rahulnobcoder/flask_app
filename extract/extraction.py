import PyPDF2
import docx
import re
import pandas as pd
import os
from spire.doc import *
from spire.doc.common import *

def extract_text_from_pdf(pdf_path):
    text = ''
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        num_pages = len(pdf_reader.pages)
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def extract_text_from_doc(doc_path):
    document = Document()
    document.LoadFromFile(doc_path)
    document_text = document.GetText()
    document.Close()
    return document_text

def extract_email_addresses(text):
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\b'
    email_pattern2 = r'\b[A-Za-z0-9._%+-]+ ?@[A-Za-z0-9.-]+ ?\.[A-Z|a-z]{2,}\b'
    return re.findall(email_pattern,text) or re.findall(email_pattern2,text)

def extract_phone_numbers(text):
    phone_pattern = r'\b\d{10}\b'
    phone_pattern0 = r'\b\d{5}\s\d{5}\b'
    phone_pattern1 = r'\+\d{2}-\d{5}-\d{5}'
    phone_pattern2 = r'\+\d{2}\s?\d{3}-\d{3}-\d{4}'
    return re.findall(phone_pattern,text) or  re.findall(phone_pattern0,text) or re.findall(phone_pattern1,text) or re.findall(phone_pattern2,text) 

def get_df(path):
    files=os.listdir(path)
    file_names=[file.split('.') for file in files]
    text=[]
    for name,ext in file_names:
        file='.'.join([name,ext])
        file_path=os.path.join(path, file) 
        if ext=='pdf':
            tx=extract_text_from_pdf(file_path)
        elif ext=='docx':
            tx=extract_text_from_docx(file_path)
        elif ext=='doc':
            tx=extract_text_from_doc(file_path)
        text.append(tx)
    emails=[]
    ph_no=[]
    for doc in text:
        email=extract_email_addresses(doc)
        if len(email)>0:
            emails.append(email[0])
        else:
            emails.append([])
        pn=extract_phone_numbers(doc)
        if len(pn)>0:
            ph_no.append(pn[0])
        else:
            ph_no.append([])
    data={'Email':emails,
          'phone_number':ph_no,
          'text':text}
    df = pd.DataFrame(data)
    return df

def clear_folder_contents(folder_path):
    # Check if the path exists and is a directory
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        # Iterate over all files and subdirectories in the folder
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)
            # Check if the item is a file
            if os.path.isfile(item_path):
                # Remove the file
                os.remove(item_path)
                print(f"File '{item_path}' removed.")
            # Check if the item is a directory
            elif os.path.isdir(item_path):
                # Remove the subdirectory and its contents recursively
                shutil.rmtree(item_path)
                print(f"Folder '{item_path}' and its contents removed.")
        print(f"All contents of folder '{folder_path}' cleared successfully.")
    else:
        print(f"Folder '{folder_path}' does not exist or is not a directory.")
